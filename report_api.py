from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import json
import os
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
import io
import uuid
import requests
from PIL import Image as PILImage, ImageDraw, ImageFont
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import numpy as np
import base64

app = Flask(__name__)
CORS(app)

# Configuration
REPORTS_FILE = 'report_requests.json'
ADMIN_EMAIL = 'your-email@gmail.com'  # Change this to your email
ADMIN_PASSWORD = 'your-app-password'  # Change this to your app password

def load_requests():
    """Load existing requests from file"""
    if os.path.exists(REPORTS_FILE):
        with open(REPORTS_FILE, 'r') as f:
            return json.load(f)
    return []

def save_requests(requests):
    """Save requests to file"""
    with open(REPORTS_FILE, 'w') as f:
        json.dump(requests, f, indent=2)

def generate_chart_image(topic, chart_type="bar"):
    """Generate a chart image based on the topic"""
    try:
        # Create figure and axis
        fig, ax = plt.subplots(figsize=(10, 6))
        
        if chart_type == "bar":
            # Generate sample data for bar chart
            categories = ['Category A', 'Category B', 'Category C', 'Category D', 'Category E']
            values = np.random.randint(20, 100, len(categories))
            
            bars = ax.bar(categories, values, color=['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7'])
            ax.set_title(f'Analysis Results for {topic}', fontsize=16, fontweight='bold', pad=20)
            ax.set_xlabel('Categories', fontsize=12)
            ax.set_ylabel('Values', fontsize=12)
            
            # Add value labels on bars
            for bar, value in zip(bars, values):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                       f'{value}', ha='center', va='bottom', fontweight='bold')
        
        elif chart_type == "pie":
            # Generate sample data for pie chart
            labels = ['Component A', 'Component B', 'Component C', 'Component D']
            sizes = np.random.randint(15, 40, len(labels))
            
            wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.1f%%',
                                             colors=['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4'])
            ax.set_title(f'Distribution Analysis for {topic}', fontsize=16, fontweight='bold', pad=20)
        
        elif chart_type == "line":
            # Generate sample data for line chart
            x = np.linspace(0, 10, 20)
            y = np.sin(x) * 50 + 50 + np.random.normal(0, 5, 20)
            
            ax.plot(x, y, marker='o', linewidth=2, markersize=6, color='#FF6B6B')
            ax.set_title(f'Trend Analysis for {topic}', fontsize=16, fontweight='bold', pad=20)
            ax.set_xlabel('Time Period', fontsize=12)
            ax.set_ylabel('Performance Metrics', fontsize=12)
            ax.grid(True, alpha=0.3)
        
        # Style the chart
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()
        
        # Save to buffer
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()
        
        return img_buffer
        
    except Exception as e:
        print(f"Error generating chart: {e}")
        return None

def generate_diagram_image(topic):
    """Generate a diagram image based on the topic"""
    try:
        # Create a new image with white background
        img = PILImage.new('RGB', (800, 600), color='white')
        draw = ImageDraw.Draw(img)
        
        # Try to use a font, fallback to default if not available
        try:
            font = ImageFont.truetype("arial.ttf", 20)
            title_font = ImageFont.truetype("arial.ttf", 24)
        except:
            font = ImageFont.load_default()
            title_font = ImageFont.load_default()
        
        # Draw title
        title = f"System Architecture: {topic}"
        draw.text((400, 30), title, fill='black', font=title_font, anchor='mm')
        
        # Draw boxes and connections
        boxes = [
            (100, 150, 300, 200, "Input Layer"),
            (500, 150, 700, 200, "Output Layer"),
            (300, 300, 500, 350, "Processing Layer"),
            (200, 400, 400, 450, "Storage Layer"),
            (400, 400, 600, 450, "Analysis Layer")
        ]
        
        for x1, y1, x2, y2, text in boxes:
            # Draw box
            draw.rectangle([x1, y1, x2, y2], outline='black', width=2, fill='lightblue')
            # Draw text
            draw.text(((x1 + x2) // 2, (y1 + y2) // 2), text, fill='black', font=font, anchor='mm')
        
        # Draw connections
        connections = [
            ((200, 175), (500, 175)),  # Input to Output
            ((200, 175), (400, 300)),  # Input to Processing
            ((400, 325), (500, 175)),  # Processing to Output
            ((400, 325), (300, 400)),  # Processing to Storage
            ((400, 325), (500, 400)),  # Processing to Analysis
        ]
        
        for start, end in connections:
            draw.line([start, end], fill='red', width=2)
        
        # Save to buffer
        img_buffer = io.BytesIO()
        img.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        return img_buffer
        
    except Exception as e:
        print(f"Error generating diagram: {e}")
        return None

def create_image_placeholder(text, width=400, height=300):
    """Create a placeholder image with text"""
    try:
        img = PILImage.new('RGB', (width, height), color='#f0f0f0')
        draw = ImageDraw.Draw(img)
        
        try:
            font = ImageFont.truetype("arial.ttf", 16)
        except:
            font = ImageFont.load_default()
        
        # Draw border
        draw.rectangle([0, 0, width-1, height-1], outline='#cccccc', width=2)
        
        # Draw text
        draw.text((width//2, height//2), text, fill='#666666', font=font, anchor='mm')
        
        # Save to buffer
        img_buffer = io.BytesIO()
        img.save(img_buffer, format='PNG')
        img_buffer.seek(0)
        
        return img_buffer
        
    except Exception as e:
        print(f"Error creating placeholder: {e}")
        return None

def send_email_notification(form_data):
    """Send email notification to admin about new report request"""
    try:
        # Email content
        subject = f"New AI Report Request: {form_data['topic']}"
        
        body = f"""
        New AI Report Generation Request
        
        Report Details:
        - Topic: {form_data['topic']}
        - Pages: {form_data['pageCount']}
        
        Requirements:
        {form_data.get('requirements', 'None specified')}
        
        Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        """
        
        # Create message
        msg = MIMEMultipart()
        msg['From'] = ADMIN_EMAIL
        msg['To'] = ADMIN_EMAIL
        msg['Subject'] = subject
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Send email (uncomment if you have email credentials)
        # server = smtplib.SMTP('smtp.gmail.com', 587)
        # server.starttls()
        # server.login(ADMIN_EMAIL, ADMIN_PASSWORD)
        # server.send_message(msg)
        # server.quit()
        
        print(f"Email notification would be sent for: {form_data['topic']}")
        return True
        
    except Exception as e:
        print(f"Error sending email: {e}")
        return False

@app.route('/api/submit-report-request', methods=['POST'])
def submit_report_request():
    """Handle AI report generation requests"""
    try:
        data = request.json
        
        # Validate required fields
        required_fields = ['topic', 'pageCount', 'requirements']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'success': False, 'error': f'Missing required field: {field}'}), 400
        
        # Generate a unique ID for this request
        request_id = str(uuid.uuid4())
        
        # Store the data temporarily in memory (not in file)
        # This will be lost when server restarts, but that's what you want
        
        return jsonify({
            'success': True,
            'message': 'Report generated successfully!',
            'request_id': request_id,
            'report_data': {
                'topic': data['topic'],
                'pageCount': data['pageCount'],
                'requirements': data['requirements']
            }
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/report-requests', methods=['GET'])
def get_report_requests():
    """Get all report requests (for admin dashboard)"""
    try:
        requests = load_requests()
        return jsonify({
            'success': True,
            'requests': requests
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/update-request-status', methods=['POST'])
def update_request_status():
    """Update request status (for admin)"""
    try:
        data = request.json
        request_id = data.get('id')
        new_status = data.get('status')
        
        if not request_id or not new_status:
            return jsonify({'success': False, 'error': 'Missing id or status'}), 400
        
        requests = load_requests()
        
        for req in requests:
            if req['id'] == request_id:
                req['status'] = new_status
                req['updated_at'] = datetime.now().isoformat()
                break
        
        save_requests(requests)
        
        return jsonify({'success': True, 'message': 'Status updated successfully'})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def generate_report_content(topic, requirements, page_count):
    """Generate AI-powered report content"""
    # This is a simplified content generator
    # In a real implementation, you would integrate with an AI service like OpenAI
    
    # Calculate target words based on page count (roughly 250-300 words per page)
    target_words = int(page_count) * 275
    
    # Base content template
    base_content = f"""
# {topic}

## Executive Summary
This report provides a comprehensive analysis of {topic}. The research covers various aspects including methodology, findings, and recommendations. The study aims to provide detailed insights into the subject matter while addressing the specific requirements outlined in the research objectives.

## Introduction
{requirements}

The introduction section establishes the context and background for this research on {topic}. It provides a foundation for understanding the significance of the study and its potential impact on the field. The research objectives are clearly defined, and the scope of the investigation is outlined to ensure comprehensive coverage of the subject matter.

## Methodology
The research methodology employed in this study includes:
- Literature review and analysis of existing research
- Data collection and analysis using appropriate tools
- Statistical evaluation and validation of findings
- Comparative study with similar research in the field
- Qualitative and quantitative analysis methods
- Peer review and expert consultation processes

## Findings
Based on the comprehensive analysis of {topic}, the following key findings emerged:
1. Primary finding related to {topic} and its implications
2. Secondary findings and their broader implications
3. Statistical significance of results and their reliability
4. Correlation between different variables and factors
5. Unexpected discoveries and their potential impact
6. Validation of initial hypotheses and research questions

## Analysis
Detailed analysis of the findings reveals:
- Critical insights about {topic} and its applications
- Comparative analysis with existing research and methodologies
- Identification of patterns, trends, and anomalies
- Statistical significance and confidence intervals
- Practical implications and real-world applications
- Limitations and areas for further research

## Recommendations
Based on the comprehensive analysis, the following recommendations are proposed:
1. Immediate actions to address key findings and implement solutions
2. Long-term strategic recommendations for sustainable development
3. Implementation guidelines and best practices
4. Resource allocation and investment priorities
5. Risk mitigation strategies and contingency plans
6. Monitoring and evaluation frameworks for ongoing assessment

## Conclusion
This report concludes that {topic} presents significant opportunities and challenges that require careful consideration and strategic planning. The findings provide a solid foundation for future research and practical applications in the field. The comprehensive analysis undertaken in this study contributes valuable insights to the existing body of knowledge and offers practical guidance for stakeholders and decision-makers.

## References
1. Author, A. (2024). Title of the study. Journal Name, Volume(Issue), Pages.
2. Researcher, B. (2023). Related research on {topic}. Publication Name.
3. Expert, C. (2024). Current trends in the field. Conference Proceedings.
4. Scholar, D. (2023). Advanced methodologies in research. Academic Press.
5. Analyst, E. (2024). Statistical analysis and interpretation. Research Quarterly.
"""
    
    # Calculate current word count
    current_words = len(base_content.split())
    
    # Generate additional content sections based on page count
    additional_sections = []
    
    if int(page_count) >= 3:
        additional_sections.append(f"""
## Technical Implementation Details

### System Architecture
The technical implementation of {topic} involves several key components that work together to create a robust and scalable solution. The system architecture is designed with modularity and flexibility in mind, allowing for easy maintenance and future enhancements.

### Technology Stack
The technology stack utilized in this project includes:
- Modern programming languages and frameworks
- Database management systems and data storage solutions
- Cloud computing platforms and infrastructure
- Security protocols and authentication mechanisms
- Performance monitoring and optimization tools

### Performance Optimization
Performance optimization strategies implemented include:
- Efficient algorithms and data structures
- Caching mechanisms and memory management
- Load balancing and scalability considerations
- Database query optimization and indexing
- Network optimization and bandwidth management
""")
    
    if int(page_count) >= 4:
        additional_sections.append(f"""
## Implementation Challenges and Solutions

### Technical Complexity
During the development and implementation of {topic}, several challenges were encountered that required innovative solutions and careful planning. The technical complexity of the project necessitated a systematic approach to problem-solving.

### Resource Management
Resource constraints and timeline management presented significant challenges:
- Budget limitations and cost optimization strategies
- Human resource allocation and skill development
- Time management and project scheduling
- Quality assurance and testing requirements
- Risk assessment and mitigation planning

### User Experience Considerations
User acceptance and adoption factors were carefully considered:
- User interface design and usability testing
- Accessibility requirements and compliance
- Training and documentation needs
- Feedback collection and iterative improvement
- Change management and organizational impact
""")
    
    if int(page_count) >= 5:
        additional_sections.append(f"""
## Future Directions and Opportunities

### Emerging Technologies
The future development of {topic} presents several exciting opportunities:
- Artificial intelligence and machine learning integration
- Blockchain technology and decentralized solutions
- Internet of Things (IoT) and connected systems
- Augmented and virtual reality applications
- Quantum computing and advanced algorithms

### Market Trends
Current market trends and industry developments include:
- Digital transformation and automation initiatives
- Sustainability and environmental considerations
- Globalization and international collaboration
- Regulatory compliance and governance frameworks
- Innovation and competitive advantage strategies

### Research Opportunities
Academic collaboration and research opportunities include:
- Interdisciplinary research partnerships
- International collaboration and knowledge sharing
- Industry-academia partnerships and joint projects
- Funding opportunities and grant applications
- Publication and dissemination strategies
""")
    
    if int(page_count) >= 6:
        additional_sections.append(f"""
## Case Studies and Practical Applications

### Case Study 1: Educational Implementation
A comprehensive case study was conducted on the implementation of {topic} in educational institutions:
- Initial assessment and needs analysis
- Implementation strategy and timeline
- Training and capacity building programs
- Evaluation and impact assessment
- Lessons learned and best practices

### Case Study 2: Commercial Applications
Commercial applications and success stories were analyzed:
- Market entry strategies and competitive positioning
- Customer acquisition and retention strategies
- Revenue generation and business model development
- Partnership and collaboration opportunities
- Scaling and expansion considerations

### Case Study 3: International Perspectives
Global impact and international perspectives were examined:
- Cultural considerations and localization requirements
- Regulatory compliance and legal frameworks
- Cross-border collaboration and partnerships
- Technology transfer and knowledge sharing
- Sustainable development and social impact
""")
    
    if int(page_count) >= 7:
        additional_sections.append(f"""
## Statistical Analysis and Validation

### Data Collection Methodology
Comprehensive statistical analysis was conducted to validate the findings:
- Sample size determination and statistical power analysis
- Data collection methods and quality assurance
- Survey design and questionnaire development
- Interview protocols and qualitative data collection
- Secondary data sources and literature review

### Statistical Tests and Analysis
Advanced statistical methods were employed:
- Descriptive statistics and data visualization
- Inferential statistics and hypothesis testing
- Correlation analysis and regression modeling
- Factor analysis and dimensionality reduction
- Time series analysis and trend identification

### Validation and Reliability
Ensuring the reliability and validity of findings:
- Internal consistency and reliability testing
- Construct validity and measurement accuracy
- External validity and generalizability assessment
- Peer review and expert validation
- Replication studies and cross-validation
""")
    
    if int(page_count) >= 8:
        additional_sections.append(f"""
## Comparative Analysis and Benchmarking

### Competitive Analysis
A comprehensive comparative analysis was performed:
- Feature comparison and competitive positioning
- Performance benchmarking and evaluation metrics
- Cost-benefit analysis and return on investment
- User satisfaction and feedback analysis
- Market share and competitive advantage assessment

### Industry Benchmarking
Industry standards and best practices were evaluated:
- Performance metrics and key performance indicators
- Quality standards and certification requirements
- Innovation and technology adoption rates
- Customer satisfaction and loyalty metrics
- Operational efficiency and productivity measures

### Gap Analysis
Identification of gaps and improvement opportunities:
- Current state assessment and baseline measurement
- Target state definition and goal setting
- Gap identification and prioritization
- Action planning and implementation roadmap
- Progress monitoring and continuous improvement
""")
    
    # Add additional sections to reach target word count
    content = base_content
    sections_added = 0
    
    for section in additional_sections:
        content += section
        sections_added += 1
        current_words = len(content.split())
        if current_words >= target_words:
            break
    
    # If still not enough content, add unique additional sections
    section_counter = 1
    while current_words < target_words and section_counter <= 5:  # Limit to prevent infinite loops
        content += f"""
## Extended Analysis Section {section_counter}

### Comprehensive Review
This section provides a comprehensive review of {topic} from multiple perspectives, including technical, economic, social, and environmental considerations. The analysis delves deeper into specific aspects that require additional examination and discussion.

### Advanced Methodologies
Advanced methodologies and analytical approaches are explored in this section, including:
- Advanced statistical modeling and predictive analytics
- Machine learning algorithms and artificial intelligence applications
- Data mining techniques and pattern recognition
- Simulation and modeling approaches
- Expert systems and decision support frameworks

### Stakeholder Perspectives
Various stakeholder perspectives are examined in detail:
- End-user requirements and expectations
- Management and organizational considerations
- Technical and operational requirements
- Regulatory and compliance perspectives
- Market and competitive landscape analysis

### Implementation Strategies
Detailed implementation strategies and best practices are discussed:
- Project planning and management methodologies
- Risk assessment and mitigation strategies
- Quality assurance and testing protocols
- Change management and organizational development
- Performance monitoring and continuous improvement
"""
        current_words = len(content.split())
        section_counter += 1
    
    return content

@app.route('/api/generate-report/<request_id>/<format>', methods=['GET'])
def generate_report(request_id, format):
    """Generate and download report in specified format"""
    try:
        # Get the topic and requirements from query parameters
        topic = request.args.get('topic', 'Sample Report Topic')
        requirements = request.args.get('requirements', 'Sample requirements for the report')
        page_count = request.args.get('pageCount', '10')
        
        # Generate report content based on actual user input
        content = generate_report_content(topic, requirements, page_count)
        
        if format.lower() == 'pdf':
            return generate_pdf_report(content, topic)
        elif format.lower() == 'docx':
            return generate_docx_report(content, topic)
        else:
            return jsonify({'success': False, 'error': 'Unsupported format'}), 400
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def generate_pdf_report(content, topic):
    """Generate PDF report with images and proper alignment"""
    try:
        # Create a buffer to store the PDF
        buffer = io.BytesIO()
        
        # Create the PDF object
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20,
            alignment=TA_LEFT
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading3'],
            fontSize=12,
            spaceAfter=8,
            spaceBefore=12,
            alignment=TA_LEFT
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            alignment=TA_JUSTIFY
        )
        
        # Parse content and create PDF elements
        elements = []
        
        # Split content into sections
        sections = content.split('\n## ')
        
        for i, section in enumerate(sections):
            if i == 0:  # First section (title)
                lines = section.split('\n')
                title = lines[0].replace('# ', '')
                elements.append(Paragraph(title, title_style))
                elements.append(Spacer(1, 20))
                
                # Add remaining content from first section
                if len(lines) > 1:
                    for line in lines[1:]:
                        if line.strip():
                            if line.startswith('### '):
                                elements.append(Paragraph(line.replace('### ', ''), subheading_style))
                            else:
                                elements.append(Paragraph(line, normal_style))
                            elements.append(Spacer(1, 6))
            else:
                # Handle other sections
                lines = section.split('\n')
                if lines[0].strip():
                    elements.append(Paragraph(lines[0], heading_style))
                    elements.append(Spacer(1, 12))
                
                # Add images for certain sections
                if "Technical Implementation" in lines[0] or "System Architecture" in lines[0]:
                    # Add diagram image
                    diagram_img = generate_diagram_image(topic)
                    if diagram_img:
                        img = Image(diagram_img, width=5*inch, height=3.75*inch)
                        img.hAlign = 'CENTER'
                        elements.append(img)
                        elements.append(Spacer(1, 12))
                
                if "Analysis" in lines[0] or "Findings" in lines[0]:
                    # Add chart image
                    chart_img = generate_chart_image(topic, "bar")
                    if chart_img:
                        img = Image(chart_img, width=5*inch, height=3*inch)
                        img.hAlign = 'CENTER'
                        elements.append(img)
                        elements.append(Spacer(1, 12))
                
                if "Statistical Analysis" in lines[0]:
                    # Add pie chart
                    pie_img = generate_chart_image(topic, "pie")
                    if pie_img:
                        img = Image(pie_img, width=4*inch, height=3*inch)
                        img.hAlign = 'CENTER'
                        elements.append(img)
                        elements.append(Spacer(1, 12))
                
                if "Trend" in lines[0] or "Future" in lines[0]:
                    # Add line chart
                    line_img = generate_chart_image(topic, "line")
                    if line_img:
                        img = Image(line_img, width=5*inch, height=3*inch)
                        img.hAlign = 'CENTER'
                        elements.append(img)
                        elements.append(Spacer(1, 12))
                
                for line in lines[1:]:
                    if line.strip():
                        if line.startswith('### '):
                            elements.append(Paragraph(line.replace('### ', ''), subheading_style))
                        elif line.startswith('- '):
                            elements.append(Paragraph('• ' + line[2:], normal_style))
                        else:
                            elements.append(Paragraph(line, normal_style))
                        elements.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        
        # Generate filename
        filename = f"{topic.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'PDF generation failed: {str(e)}'}), 500

def generate_docx_report(content, topic):
    """Generate DOCX report with images and proper alignment"""
    try:
        # Create a new Document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Parse content and add to document
        sections = content.split('\n## ')
        
        for i, section in enumerate(sections):
            if i == 0:  # First section (title)
                lines = section.split('\n')
                title = lines[0].replace('# ', '')
                heading = doc.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add remaining content from first section
                if len(lines) > 1:
                    for line in lines[1:]:
                        if line.strip():
                            if line.startswith('### '):
                                doc.add_heading(line.replace('### ', ''), 2)
                            else:
                                p = doc.add_paragraph(line)
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                # Handle other sections
                lines = section.split('\n')
                if lines[0].strip():
                    doc.add_heading(lines[0], 1)
                
                # Add images for certain sections
                if "Technical Implementation" in lines[0] or "System Architecture" in lines[0]:
                    # Add diagram image
                    diagram_img = generate_diagram_image(topic)
                    if diagram_img:
                        # Save image temporarily and add to document
                        img_buffer = io.BytesIO()
                        PILImage.open(diagram_img).save(img_buffer, format='PNG')
                        img_buffer.seek(0)
                        doc.add_picture(img_buffer, width=Inches(5))
                        # Center the image
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()  # Add space after image
                
                if "Analysis" in lines[0] or "Findings" in lines[0]:
                    # Add chart image
                    chart_img = generate_chart_image(topic, "bar")
                    if chart_img:
                        img_buffer = io.BytesIO()
                        PILImage.open(chart_img).save(img_buffer, format='PNG')
                        img_buffer.seek(0)
                        doc.add_picture(img_buffer, width=Inches(5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                
                if "Statistical Analysis" in lines[0]:
                    # Add pie chart
                    pie_img = generate_chart_image(topic, "pie")
                    if pie_img:
                        img_buffer = io.BytesIO()
                        PILImage.open(pie_img).save(img_buffer, format='PNG')
                        img_buffer.seek(0)
                        doc.add_picture(img_buffer, width=Inches(4))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                
                if "Trend" in lines[0] or "Future" in lines[0]:
                    # Add line chart
                    line_img = generate_chart_image(topic, "line")
                    if line_img:
                        img_buffer = io.BytesIO()
                        PILImage.open(line_img).save(img_buffer, format='PNG')
                        img_buffer.seek(0)
                        doc.add_picture(img_buffer, width=Inches(5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                
                for line in lines[1:]:
                    if line.strip():
                        if line.startswith('### '):
                            doc.add_heading(line.replace('### ', ''), 2)
                        elif line.startswith('- '):
                            p = doc.add_paragraph(line[2:], style='List Bullet')
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            p = doc.add_paragraph(line)
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Generate filename
        filename = f"{topic.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'DOCX generation failed: {str(e)}'}), 500

if __name__ == '__main__':
    # For production deployment
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
